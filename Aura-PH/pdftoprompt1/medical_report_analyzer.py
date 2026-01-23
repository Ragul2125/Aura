import sys
import os
from pathlib import Path
import argparse
from typing import List
import shutil

# Try importing required libraries
try:
    import pdfplumber
except ImportError:
    print("Error: pdfplumber is not installed. Please install it using: pip install pdfplumber")
    sys.exit(1)

try:
    import requests
except ImportError:
    print("Error: requests is not installed. Please install it using: pip install requests")
    sys.exit(1)

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("Warning: OCR libraries not available. Image-based PDFs and image files will not be processed.")
    print("Install with: pip install pytesseract pillow pdf2image")

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Word documents (.docx) will not be supported.")
    print("Install with: pip install python-docx")


def extract_text_with_ocr(image_path: str) -> str:
    """
    Extract text from an image using OCR.
    
    Args:
        image_path: Path to the image file
        
    Returns:
        Extracted text as a string
    """
    if not OCR_AVAILABLE:
        raise Exception("OCR is not available. Please install pytesseract and pillow.")
    
    try:
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        raise Exception(f"Error performing OCR: {str(e)}")


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text from a PDF file. Tries text extraction first, then OCR if needed.
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        Extracted text as a string
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")
    
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            num_pages = len(pdf.pages)
            
            print(f"Processing PDF with {num_pages} page(s)...")
            
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text and len(page_text.strip()) > 10:  # Meaningful text found
                    text += f"\n--- Page {page_num} ---\n"
                    text += page_text
                else:
                    # Try extracting tables if no text found
                    tables = page.extract_tables()
                    if tables:
                        text += f"\n--- Page {page_num} (Tables) ---\n"
                        for table in tables:
                            for row in table:
                                if row:
                                    text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
                
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")
    
    # If no text extracted and OCR is available, try OCR
    if not text.strip():
        if not OCR_AVAILABLE:
            raise ValueError(
                "No text could be extracted from the PDF (likely scanned/image-based). "
                "OCR support is not installed. Install with: pip install pytesseract pillow pdf2image"
            )
        if not PDF2IMAGE_AVAILABLE:
            raise ValueError(
                "No text could be extracted from the PDF (likely scanned/image-based). "
                "OCR requires pdf2image. Install with: pip install pdf2image"
            )

        # pdf2image requires Poppler (pdftoppm/pdfinfo) to be installed and available on PATH
        if shutil.which("pdftoppm") is None or shutil.which("pdfinfo") is None:
            raise ValueError(
                "OCR requires Poppler, but it was not found on your PATH.\n"
                "- macOS: brew install poppler\n"
                "- Ubuntu/Debian: sudo apt-get install poppler-utils\n"
                "- Windows: install Poppler and add its bin folder to PATH\n"
            )

        print("No text found in PDF. Attempting OCR on image-based PDF...")
        try:
            images = convert_from_path(pdf_path)
            for page_num, image in enumerate(images, 1):
                print(f"Performing OCR on page {page_num}...")
                page_text = pytesseract.image_to_string(image)
                if page_text.strip():
                    text += f"\n--- Page {page_num} (OCR) ---\n"
                    text += page_text
        except Exception as e:
            raise Exception(
                "OCR failed while converting PDF pages to images. "
                "If you see 'Unable to get page count', install Poppler.\n"
                f"Details: {str(e)}"
            )
    
    if not text.strip():
        raise ValueError(
            "No text could be extracted from the PDF. It may be corrupted, encrypted, or OCR failed."
        )
    
    return text


def extract_text_from_image(image_path: str) -> str:
    """
    Extract text from an image file (JPEG, PNG) using OCR.
    
    Args:
        image_path: Path to the image file
        
    Returns:
        Extracted text as a string
    """
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"Image file not found: {image_path}")
    
    if not OCR_AVAILABLE:
        raise Exception("OCR is not available. Please install pytesseract and pillow.")
    
    print("Performing OCR on image...")
    text = extract_text_with_ocr(image_path)
    
    if not text.strip():
        raise ValueError("No text could be extracted from the image.")
    
    return text


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from a Word document (.docx).
    
    Args:
        docx_path: Path to the Word document
        
    Returns:
        Extracted text as a string
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Word document not found: {docx_path}")
    
    if not DOCX_AVAILABLE:
        raise Exception("python-docx is not available. Please install it using: pip install python-docx")
    
    try:
        doc = Document(docx_path)
        text = ""
        for para_num, paragraph in enumerate(doc.paragraphs, 1):
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table_num, table in enumerate(doc.tables, 1):
            text += f"\n--- Table {table_num} ---\n"
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    text += row_text + "\n"
        
        if not text.strip():
            raise ValueError("No text could be extracted from the Word document.")
        
        return text
    except Exception as e:
        raise Exception(f"Error reading Word document: {str(e)}")


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from a file (PDF, JPEG, PNG, or Word document).
    Automatically detects file type and uses appropriate extraction method.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text as a string
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']:
        return extract_text_from_image(file_path)
    elif file_ext in ['.docx', '.doc']:
        if file_ext == '.doc':
            raise ValueError("Old .doc format is not supported. Please convert to .docx format.")
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: PDF, JPEG, PNG, DOCX")


def chunk_text(text: str, chunk_size: int = 3000, overlap: int = 200) -> List[str]:
    """
    Split text into chunks with overlap to maintain context.
    
    Args:
        text: The text to chunk
        chunk_size: Maximum size of each chunk
        overlap: Number of characters to overlap between chunks
        
    Returns:
        List of text chunks
    """
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Try to break at a sentence or paragraph boundary
        if end < len(text):
            # Look for sentence endings
            for break_char in ['. ', '\n\n', '\n']:
                last_break = text.rfind(break_char, start, end)
                if last_break != -1:
                    end = last_break + len(break_char)
                    break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap  # Overlap for context
    
    return chunks


def analyze_with_llm(text_chunks: List[str], api_key: str, base_url: str = None, model: str = None) -> str:
    """
    Send text chunks to NVIDIA LLM and get a brief medical summary.
    
    Args:
        text_chunks: List of text chunks from the PDF
        api_key: NVIDIA API key
        base_url: NVIDIA API base URL (optional, defaults to NIM endpoint)
        model: Model name (optional, defaults to llama-3.1-8b-instruct)
        
    Returns:
        Brief medical summary
    """
    # Combine all chunks for analysis
    full_text = "\n\n".join(text_chunks)
    
    # Truncate if too long (NVIDIA API has token limits)
    max_chars = 12000  # Roughly 3000 tokens
    if len(full_text) > max_chars:
        print(f"Warning: Text is very long ({len(full_text)} chars). Truncating to {max_chars} chars for analysis.")
        full_text = full_text[:max_chars] + "\n\n[... text truncated ...]"
    
    prompt = """You are a medical assistant. Analyze the following medical report and provide a brief, concise summary of the key findings.

Focus on:
- Abnormal values (high/low indicators)
- Important medical conditions or concerns
- Key test results that need attention

Format your response as a brief, easy-to-understand summary. Use simple language like:
"Cholesterol is high, blood pressure is low, etc."

Medical Report:
{}

Brief Summary:""".format(full_text)
    
    # Default NVIDIA NIM endpoint
    if base_url is None:
        base_url = "https://integrate.api.nvidia.com/v1"
    
    # Default model
    if model is None:
        model = "meta/llama-3.1-8b-instruct"
    
    # NVIDIA API endpoint
    api_url = f"{base_url}/chat/completions"
    
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "You are a helpful medical assistant that provides brief, clear summaries of medical reports."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3,
        "max_tokens": 500,
        "top_p": 0.7,
        "stream": False
    }
    
    try:
        print("Sending to NVIDIA LLM for analysis...")
        response = requests.post(
            api_url,
            json=payload,
            headers=headers,
            timeout=60
        )
        
        response.raise_for_status()
        result = response.json()
        
        # Extract the summary from NVIDIA API response
        if "choices" in result and len(result["choices"]) > 0:
            summary = result["choices"][0]["message"]["content"].strip()
            return summary
        else:
            raise Exception("Unexpected response format from NVIDIA API")
        
    except requests.exceptions.HTTPError as e:
        error_msg = f"HTTP Error: {e.response.status_code}"
        try:
            error_detail = e.response.json()
            error_msg += f" - {error_detail}"
        except:
            error_msg += f" - {e.response.text}"
        raise Exception(f"Error calling NVIDIA API: {error_msg}")
    except requests.exceptions.RequestException as e:
        raise Exception(f"Error calling NVIDIA API: {str(e)}")


def main():
    parser = argparse.ArgumentParser(
        description="Analyze a medical report PDF and generate a brief summary",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Example usage:
  python medical_report_analyzer.py /path/to/report.pdf
  python medical_report_analyzer.py /path/to/report.png
  python medical_report_analyzer.py /path/to/report.docx
  
Supported formats: PDF, JPEG, PNG, DOCX
  
Environment variable:
  NVIDIA_API_KEY - Your NVIDIA API key (required)
  
Available models:
  - meta/llama-3.1-8b-instruct (default)
  - meta/llama-3.1-70b-instruct
  - mistralai/mistral-7b-instruct-v0.2
  - google/gemma-2-9b-it
        """
    )
    parser.add_argument(
        "file_path",
        type=str,
        help="Path to the medical report file (PDF, JPEG, PNG, or Word document)"
    )
    parser.add_argument(
        "--api-key",
        type=str,
        default=None,
        help="NVIDIA API key (or set NVIDIA_API_KEY environment variable)"
    )
    parser.add_argument(
        "--base-url",
        type=str,
        default=None,
        help="NVIDIA API base URL (default: https://integrate.api.nvidia.com/v1)"
    )
    parser.add_argument(
        "--model",
        type=str,
        default="meta/llama-3.1-8b-instruct",
        help="NVIDIA model to use (default: meta/llama-3.1-8b-instruct). Examples: meta/llama-3.1-70b-instruct, mistralai/mistral-7b-instruct-v0.2"
    )
    
    args = parser.parse_args()
    
    # Get API key
    api_key = args.api_key or os.getenv("NVIDIA_API_KEY")
    if not api_key:
        print("Error: NVIDIA API key is required.")
        print("Please set it as an environment variable: export NVIDIA_API_KEY='your-key'")
        print("Or pass it with --api-key flag")
        print("\nGet your API key from: https://build.nvidia.com/")
        sys.exit(1)
    
    try:
        # Extract text from file
        print(f"\n{'='*60}")
        print(f"Medical Report Analyzer")
        print(f"{'='*60}\n")
        print(f"File: {args.file_path}\n")
        
        text = extract_text_from_file(args.file_path)
        print(f"Extracted {len(text)} characters from file\n")
        
        # Chunk the text
        chunks = chunk_text(text)
        print(f"Split into {len(chunks)} chunk(s) for processing\n")
        
        # Analyze with LLM
        summary = analyze_with_llm(chunks, api_key, args.base_url, args.model)
        
        # Display results
        print(f"\n{'='*60}")
        print("MEDICAL REPORT SUMMARY")
        print(f"{'='*60}\n")
        print(summary)
        print(f"\n{'='*60}\n")
        
    except FileNotFoundError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    except ValueError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
